"""Lightweight Markdown to .docx converter for the BIN_GUIDED_TILE_PYRAMID
runbook. Handles the specific MD features used in that doc:
  - H1/H2/H3/H4 headings (#, ##, ###, ####)
  - Bold **text** and italic *text*
  - Inline code `text`
  - Fenced code blocks ```...```
  - Pipe tables | col1 | col2 |
  - Bullet lists (- ...)
  - Numbered lists (1. ...)
  - Horizontal rules (---)
  - Plain paragraphs

Not a general-purpose converter. Pragmatic, tuned for the input doc.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_inline_runs(paragraph, text, base_font=None):
    """Parse a line of text for **bold**, *italic*, `code` and add runs."""
    # Split on the inline markers — keep delimiters so we know what they were.
    # Regex captures: bold, italic, inline code as separate alternatives.
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)
        if base_font:
            for run in paragraph.runs:
                run.font.name = base_font


def parse_table_row(line):
    """Parse a pipe-delimited table row into cells."""
    # Strip leading/trailing pipes, split on internal pipes
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line):
    """Returns True for lines like |---|---|---|."""
    line = line.strip()
    if not line.startswith("|"):
        return False
    cells = parse_table_row(line)
    return all(re.match(r'^:?-+:?$', cell) for cell in cells if cell)


def set_cell_background(cell, hex_color):
    """Set the background color of a docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_code_block(doc, lines):
    """Add a fenced code block as a single styled paragraph."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    # Light gray box around it would need XML shading on the paragraph;
    # for simplicity, just the monospace font + indent.


def add_table(doc, rows):
    """Add a markdown-style table to the doc."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Normalize all rows to n_cols
    rows = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            # Header row gets bold + background
            if i == 0:
                set_cell_background(cell, "D9E2F3")
            # Clear default paragraph and add our own
            cell.paragraphs[0].text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text)
            if i == 0:
                for run in p.runs:
                    run.bold = True


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    # Set default font and margins
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Narrow margins for more content per page
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Heading style tweaks
    for level in range(1, 5):
        h = doc.styles[f"Heading {level}"]
        h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            j = i + 1
            block = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                block.append(lines[j])
                j += 1
            add_code_block(doc, block)
            i = j + 1
            continue

        # Horizontal rule
        if stripped in ("---", "***"):
            p = doc.add_paragraph()
            run = p.add_run("─" * 60)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            i += 1
            continue

        # Headings
        if stripped.startswith("####"):
            doc.add_heading(stripped[4:].strip(), level=4)
            i += 1
            continue
        if stripped.startswith("###"):
            doc.add_heading(stripped[3:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith("##"):
            doc.add_heading(stripped[2:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#"):
            doc.add_heading(stripped[1:].strip(), level=1)
            i += 1
            continue

        # Table: peek-ahead for separator
        if stripped.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            header = parse_table_row(line)
            rows = [header]
            j = i + 2
            while j < len(lines) and lines[j].strip().startswith("|"):
                rows.append(parse_table_row(lines[j]))
                j += 1
            add_table(doc, rows)
            doc.add_paragraph()  # spacer
            i = j
            continue

        # Bullet list
        if re.match(r'^\s*[-*]\s', line):
            p = doc.add_paragraph(style="List Bullet")
            indent_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
            content = indent_match.group(2)
            add_inline_runs(p, content)
            i += 1
            continue

        # Numbered list
        if re.match(r'^\s*\d+\.\s', line):
            p = doc.add_paragraph(style="List Number")
            content = re.match(r'^\s*\d+\.\s+(.*)', line).group(1)
            add_inline_runs(p, content)
            i += 1
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))
    return docx_path


def main():
    if len(sys.argv) != 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    md_path = Path(sys.argv[1])
    docx_path = Path(sys.argv[2])
    if not md_path.exists():
        print(f"ERROR: {md_path} not found")
        sys.exit(1)
    convert(md_path, docx_path)
    size_mb = docx_path.stat().st_size / 1024
    print(f"wrote {docx_path} ({size_mb:.1f} KB)")


if __name__ == "__main__":
    main()
